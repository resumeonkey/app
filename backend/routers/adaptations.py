"""
Adaptations: create, list, get, trigger processing.
"""
import os
import re
import tempfile
from datetime import datetime, timezone
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from pydantic import BaseModel
from sqlalchemy.orm import Session

from backend.config import get_settings
from backend.database import get_db
from backend.models.master import MasterResume
from backend.models.adaptation import Adaptation
from backend.models.context import UserContext
from backend.services.adapter import run_adaptation
from backend.services.docx_builder import build_adapted_docx
from backend.services import storage

router   = APIRouter()
settings = get_settings()


class CreateAdaptationRequest(BaseModel):
    job_description: str
    user_instructions: Optional[str] = None
    llm_provider: str = "openai"
    llm_model: str = "gpt-4o"
    job_url: Optional[str] = None   # URL of the job listing (for tracking)


class AdaptationSummary(BaseModel):
    id: str
    master_id: str
    job_title: Optional[str]
    company_name: Optional[str]
    status: str
    created_at: datetime
    sections_changed: list[str]

    class Config:
        from_attributes = True


class AdaptationDetail(AdaptationSummary):
    job_description: str
    user_instructions: Optional[str]
    job_analysis: Optional[dict]
    blocks_changed: Optional[list]
    llm_provider: str
    llm_model: str
    error_msg: Optional[str]


# ── Background task ───────────────────────────────────────────────────────────

async def _run_adaptation_task(adaptation_id: str):
    from backend.database import SessionLocal

    db = SessionLocal()
    try:
        adaptation = db.query(Adaptation).filter(Adaptation.id == adaptation_id).first()
        if not adaptation:
            return

        master = db.query(MasterResume).filter(MasterResume.id == adaptation.master_id).first()
        if not master:
            adaptation.status = "error"
            adaptation.error_msg = "Master resume not found."
            db.commit()
            return

        adaptation.status = "processing"
        db.commit()

        # Load active context items
        active_contexts = (
            db.query(UserContext)
            .filter(UserContext.is_active == True)
            .order_by(UserContext.created_at.asc())
            .all()
        )
        user_context = "\n\n".join(
            f"[{c.title}]\n{c.content}" for c in active_contexts
        )

        result = await run_adaptation(
            master_sections=master.sections or {},
            master_full_text=master.full_text or "",
            job_description=adaptation.job_description,
            user_instructions=adaptation.user_instructions or "",
            llm_provider=adaptation.llm_provider,
            llm_model=adaptation.llm_model,
            user_context=user_context,
        )

        adaptation.job_analysis   = result["job_analysis"]
        adaptation.blocks_changed = result["blocks_changed"]

        analysis = result["job_analysis"]
        if not adaptation.job_title:
            adaptation.job_title = analysis.get("job_title", "")
        if not adaptation.company_name:
            adaptation.company_name = analysis.get("company_name", "")

        # Build adapted DOCX — filename: CandidateName_JobTitle_Company.docx
        def _slug(s: str, maxlen: int = 22) -> str:
            return re.sub(r'[^\w]', '_', s).strip('_')[:maxlen]

        safe_name    = _slug(master.candidate_name or "Resume")
        safe_title   = _slug(adaptation.job_title or "role")
        safe_company = _slug(adaptation.company_name or "company")
        output_name  = f"{safe_name}_{safe_title}_{safe_company}.docx"
        output_tmp   = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name

        try:
            if master.file_type == "docx":
                # Preserve original formatting — patch paragraphs in-place
                master_tmp = storage.download_to_temp(master.file_path, suffix=".docx")
                try:
                    build_adapted_docx(
                        master_path=master_tmp,
                        master_sections=master.sections or {},
                        blocks_changed=result["blocks_changed"],
                        output_path=output_tmp,
                    )
                finally:
                    if storage.is_remote() and os.path.exists(master_tmp):
                        os.remove(master_tmp)
            else:
                # PDF master — build a plain DOCX from the adapted full text
                _build_docx_from_text(
                    master_full_text=master.full_text or "",
                    blocks_changed=result["blocks_changed"],
                    output_path=output_tmp,
                )

            with open(output_tmp, "rb") as f:
                output_bytes = f.read()
        finally:
            if os.path.exists(output_tmp):
                os.remove(output_tmp)

        storage_path = f"outputs/{output_name}"
        saved_path = storage.upload_file(
            output_bytes, storage_path,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        adaptation.output_path = saved_path

        adaptation.status = "done"
        db.commit()

    except Exception as exc:
        # Use the same session — avoids losing the exception context
        try:
            adaptation = db.query(Adaptation).filter(Adaptation.id == adaptation_id).first()
            if adaptation:
                adaptation.status    = "error"
                adaptation.error_msg = str(exc)
                db.commit()
        except Exception:
            pass
    finally:
        db.close()


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.post("/", response_model=AdaptationDetail, status_code=201)
async def create_adaptation(
    body: CreateAdaptationRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
):
    master = db.query(MasterResume).filter(MasterResume.is_active == True).first()
    if not master:
        raise HTTPException(400, "No active master resume. Upload one first.")

    adaptation = Adaptation(
        master_id=master.id,
        job_description=body.job_description,
        user_instructions=body.user_instructions,
        llm_provider=body.llm_provider,
        llm_model=body.llm_model,
        job_url=body.job_url,
        status="pending",
    )
    db.add(adaptation)
    db.commit()
    db.refresh(adaptation)

    background_tasks.add_task(_run_adaptation_task, adaptation.id)
    return _to_detail(adaptation)


@router.get("/", response_model=list[AdaptationSummary])
def list_adaptations(db: Session = Depends(get_db)):
    adaptations = (
        db.query(Adaptation)
        .order_by(Adaptation.created_at.desc())
        .limit(50)
        .all()
    )
    return [_to_summary(a) for a in adaptations]


@router.get("/{adaptation_id}", response_model=AdaptationDetail)
def get_adaptation(adaptation_id: str, db: Session = Depends(get_db)):
    a = db.query(Adaptation).filter(Adaptation.id == adaptation_id).first()
    if not a:
        raise HTTPException(404, "Adaptation not found")
    return _to_detail(a)


@router.patch("/{adaptation_id}/applied")
def toggle_applied(adaptation_id: str, db: Session = Depends(get_db)):
    """Mark/unmark an adaptation as 'I applied to this job'."""
    a = db.query(Adaptation).filter(Adaptation.id == adaptation_id).first()
    if not a:
        raise HTTPException(404, "Adaptation not found")
    if a.applied_at:
        a.applied_at = None          # toggle off
    else:
        a.applied_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(a)
    return _to_detail(a)


@router.delete("/{adaptation_id}", status_code=204)
def delete_adaptation(adaptation_id: str, db: Session = Depends(get_db)):
    import os
    a = db.query(Adaptation).filter(Adaptation.id == adaptation_id).first()
    if not a:
        raise HTTPException(404, "Adaptation not found")
    if a.output_path:
        storage.delete_file(a.output_path)
    db.delete(a)
    db.commit()


# ── Helpers ───────────────────────────────────────────────────────────────────

def _to_summary(a: Adaptation) -> dict:
    changed = [b["section"] for b in (a.blocks_changed or [])]
    return {
        "id": a.id, "master_id": a.master_id,
        "job_title": a.job_title, "company_name": a.company_name,
        "status": a.status, "created_at": a.created_at,
        "sections_changed": changed,
        "job_url": a.job_url,
        "applied_at": a.applied_at.isoformat() if a.applied_at else None,
    }

def _to_detail(a: Adaptation) -> dict:
    return {
        **_to_summary(a),
        "job_description": a.job_description,
        "user_instructions": a.user_instructions,
        "job_analysis": a.job_analysis,
        "blocks_changed": a.blocks_changed,
        "llm_provider": a.llm_provider,
        "llm_model": a.llm_model,
        "error_msg": a.error_msg,
    }


def _repair_pdf_text(text: str) -> str:
    """
    PDFs with complex layouts (individually positioned characters, fancy headers)
    produce text where each letter lands on its own line when extracted.
    This function detects runs of single-character lines and joins them back
    into words, making the output at least readable.
    """
    lines = text.split("\n")
    output: list[str] = []
    char_buffer: list[str] = []

    def flush_buffer():
        if char_buffer:
            output.append("".join(char_buffer))
            char_buffer.clear()

    for line in lines:
        s = line.strip()
        # A "broken" line is 1-2 chars (letter, digit, punctuation fragment)
        if len(s) <= 2 and s and not s.isspace():
            char_buffer.append(s)
        else:
            flush_buffer()
            output.append(s)

    flush_buffer()
    return "\n".join(output)


def _build_docx_from_text(master_full_text: str, blocks_changed: list, output_path: str) -> None:
    """
    Build a plain DOCX from master full text (used when master is PDF).
    Replaces each original section text with the adapted version, repairs
    broken single-character lines from complex PDF layouts, then writes
    line-by-line with basic heading detection.

    NOTE: for best formatting results the user should upload a DOCX master.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.shared import Inches

    # Replace original text with adapted text in-order
    adapted_text = master_full_text
    for block in blocks_changed:
        orig = (block.get("original") or "").strip()
        adpt = (block.get("adapted") or "").strip()
        if orig and adpt:
            adapted_text = adapted_text.replace(orig, adpt, 1)

    # Repair broken single-char lines from complex PDF extraction
    adapted_text = _repair_pdf_text(adapted_text)

    doc = Document()

    # Tighter margins (1 in left/right, 0.75 in top/bottom)
    for sec in doc.sections:
        sec.top_margin    = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

    # Bullet-only lines (just "•", "-", "*", "–") with no text are artifacts
    # of broken PDF extraction — skip them entirely
    _BULLET_ONLY = re.compile(r'^[•\-\*–·]+\s*$')

    for line in adapted_text.split("\n"):
        stripped = line.strip()
        if not stripped or _BULLET_ONLY.match(stripped):
            doc.add_paragraph()  # spacing only, no visible content
            continue

        para = doc.add_paragraph()
        run  = para.add_run(stripped)
        run.font.name = "Calibri"
        run.font.size = Pt(11)

        # Heuristic: ALL-CAPS short line → section header
        if stripped.isupper() and 2 < len(stripped) < 60:
            run.bold = True
            run.font.size = Pt(12)
            # Add a thin bottom border to mimic section divider
            try:
                from docx.oxml import OxmlElement
                pPr = para._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "4")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "auto")
                pBdr.append(bottom)
                pPr.append(pBdr)
            except Exception:
                pass  # non-critical styling

    doc.save(output_path)

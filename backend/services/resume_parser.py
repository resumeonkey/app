"""
Parse a DOCX or PDF resume into a structured section map.

Output format:
{
  "full_text": "...",
  "candidate_name": "...",
  "sections": {
    "summary":    {"raw_text": "...", "para_indices": [2, 3]},
    "skills":     {"raw_text": "...", "para_indices": [5, 6, 7]},
    "experience": {"raw_text": "...", "para_indices": [9..25]},
    "education":  {"raw_text": "...", "para_indices": [26..30]},
    ...
  }
}
"""
import re
from typing import Any


# ── Section heading patterns (case-insensitive) ────────────────────────────────
SECTION_PATTERNS = {
    "summary": (
        r"(summary|profile|professional\s+profile|professional\s+summary|"
        r"career\s+profile|career\s+summary|about\s+me|objective|"
        r"sobre\s+m[ií]|perfil|resumen)"
    ),
    "skills": (
        r"(skills|technical\s+skills|core\s+competencies|key\s+skills|"
        r"core\s+capabilities|areas\s+of\s+expertise|expertise|"
        r"competencias|habilidades)"
    ),
    "experience": (
        r"(\bexperience\b|work\s+experience|professional\s+experience|"
        r"work\s+history|job\s+history|employment\s+history|\bemployment\b|"
        r"career\s+history|professional\s+background|"
        r"\bexperiencia\b|historial\s+laboral|trayectoria)"
    ),
    "education": (
        r"(education|academic|academic\s+background|"
        r"educaci[oó]n|formaci[oó]n|estudios)"
    ),
    "projects": (
        r"(projects|key\s+projects|selected\s+projects|notable\s+projects|"
        r"proyectos)"
    ),
    "certifications": (
        r"(certifications?|licenses?|credentials|"
        r"professional\s+development|training|"
        r"cursos|certificaciones?|formaci[oó]n\s+complementaria)"
    ),
    "languages": r"(languages?|idiomas?|language\s+skills)",
    "volunteer":  r"(volunteer|volunteering|community|voluntariado)",
}


def _para_meta(para, index: int, in_table: bool = False) -> dict:
    """Extract metadata dict for a single python-docx Paragraph object."""
    text = para.text.strip().replace("\n", " ")
    style = para.style.name if para.style else ""

    runs_meta: list[dict] = []
    for run in para.runs:
        run_text = run.text.replace("\n", " ")
        runs_meta.append({
            "text":    run_text,
            "bold":    bool(run.bold),
            "italic":  bool(run.italic),
            "size_pt": round(run.font.size / 12700, 1) if run.font.size else None,
        })

    has_bold     = any(r["bold"] for r in runs_meta)
    partial_bold = has_bold and not all(r["bold"] for r in runs_meta)
    _pPr = para._p.pPr
    has_numpr = (
        _pPr is not None
        and _pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr") is not None
    )

    return {
        "index":        index,
        "text":         text,
        "style":        style,
        "empty":        not text,
        "in_table":     in_table,
        "runs":         runs_meta,
        "run_count":    len(runs_meta),
        "has_bold":     has_bold,
        "partial_bold": partial_bold,
        "has_numpr":    has_numpr,
    }


def parse_docx(file_path: str) -> dict[str, Any]:
    from docx import Document
    from docx.text.paragraph import Paragraph as DocxParagraph
    from docx.oxml.ns import qn

    doc = Document(file_path)

    # ── Walk body elements in document order ─────────────────────────────────
    # doc.paragraphs only returns BODY-level paragraphs — it silently skips
    # paragraphs inside <w:tbl> elements.  Canadian resume templates frequently
    # use Word tables for the Core Competencies / Skills section, so those cells
    # are completely invisible to the old approach.
    #
    # We iterate doc.element.body children directly so that body paragraphs and
    # table-cell paragraphs are interleaved in true document order, each
    # assigned a sequential index.  docx_builder already iterates root.iter("p")
    # which includes table cells, so replacements land correctly.

    paragraphs: list[dict] = []
    idx = 0

    for child in doc.element.body:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if local == "p":
            # Body-level paragraph
            para = DocxParagraph(child, doc)
            paragraphs.append(_para_meta(para, idx, in_table=False))
            idx += 1

        elif local == "tbl":
            # Table — iterate rows → cells → paragraphs in order
            for tr in child.findall(".//" + qn("w:tr")):
                for tc in tr.findall(qn("w:tc")):
                    for p_elem in tc.findall(qn("w:p")):
                        para = DocxParagraph(p_elem, doc)
                        paragraphs.append(_para_meta(para, idx, in_table=True))
                        idx += 1

    return _build_section_map(paragraphs)


def parse_pdf(file_path: str) -> dict[str, Any]:
    lines = _extract_pdf_lines(file_path)

    # Convert to paragraph-like objects
    paragraphs = []
    for i, line in enumerate(lines):
        text = line.strip()
        # Detect likely headings: ALL CAPS or short bold-ish lines
        style = "Heading 1" if _looks_like_heading(text) else "Normal"
        paragraphs.append({"index": i, "text": text, "style": style, "empty": not text})

    return _build_section_map(paragraphs)


def _extract_pdf_lines(file_path: str) -> list[str]:
    """Extract text lines from a PDF, repairing broken single-char layouts."""
    try:
        import pdfplumber
        lines: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                lines.extend(text.split("\n"))
    except ImportError:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        lines = []
        for page in reader.pages:
            lines.extend((page.extract_text() or "").split("\n"))
    return _repair_single_char_lines(lines)


def pdf_to_docx(pdf_path: str, docx_path: str) -> None:
    """
    Convert a PDF resume into a clean, well-structured DOCX.

    This is a *content* reconstruction, not a visual one: it extracts the text
    and rebuilds it as flowing Word paragraphs with Heading styles on detected
    section titles. The original PDF's colours/columns/icons are dropped, but the
    result is a real .docx that the adaptation + export pipeline can edit
    paragraph-by-paragraph without breaking formatting.
    """
    from docx import Document

    doc = Document()
    for raw in _extract_pdf_lines(pdf_path):
        text = raw.strip()
        if not text:
            doc.add_paragraph("")
            continue
        if _looks_like_heading(text):
            doc.add_heading(text, level=1)
        else:
            doc.add_paragraph(text)
    doc.save(docx_path)


def _repair_single_char_lines(lines: list[str]) -> list[str]:
    """
    Join consecutive single/double-character lines back into words.
    Handles PDFs where complex layouts cause each letter to be
    extracted on its own line (e.g. vertically stacked name headers).
    """
    output: list[str] = []
    char_buffer: list[str] = []

    def flush():
        if char_buffer:
            output.append("".join(char_buffer))
            char_buffer.clear()

    for line in lines:
        s = line.strip()
        if len(s) <= 2 and s and not s.isspace():
            char_buffer.append(s)
        else:
            flush()
            output.append(s)

    flush()
    return output


def _looks_like_heading(text: str) -> bool:
    if not text:
        return False
    if text.isupper() and 2 < len(text) < 40:
        return True
    if re.match(r"^[A-Z][A-Z\s&/]+$", text) and len(text) < 40:
        return True
    return False


def _build_section_map(paragraphs: list[dict]) -> dict[str, Any]:
    full_lines = [p["text"] for p in paragraphs if p["text"]]
    full_text  = "\n".join(full_lines)

    # Detect candidate name: usually one of the first 3 non-empty lines
    candidate_name = _detect_name(paragraphs)

    # Find section boundaries by matching headings
    section_starts: list[tuple[str, int]] = []
    for p in paragraphs:
        if not p["text"]:
            continue
        # Table cells are CONTENT, never section headings. A competencies table
        # cell like "User Training Support" must not be mistaken for a heading
        # (it contains "training" → would false-match "certifications" and split
        # the section, scattering content across the wrong sections).
        if p.get("in_table"):
            continue
        matched = _match_section(p["text"])
        if matched:
            section_starts.append((matched, p["index"]))

    # Build sections: from one heading to the next
    # If the same section key appears more than once (e.g. "Core Capabilities"
    # and "Other Technical Skills & Tools" both match "skills"), merge them
    # so no content is lost.
    sections: dict[str, dict] = {}
    for i, (section_name, start_idx) in enumerate(section_starts):
        end_idx = section_starts[i + 1][1] if i + 1 < len(section_starts) else len(paragraphs)
        section_paras = [
            p for p in paragraphs
            if start_idx <= p["index"] < end_idx and p["text"]
        ]
        # Skip the heading line itself (first paragraph)
        content_paras = section_paras[1:] if section_paras else []
        new_text     = "\n".join(p["text"] for p in content_paras)
        new_indices  = [p["index"] for p in content_paras]
        # Per-line format profile: one entry per content paragraph, in order.
        # Each entry mirrors the python-docx run metadata captured in parse_docx.
        # docx_builder can use this as the authoritative format spec for each line.
        new_formats  = [
            {
                "text":         p["text"],
                "runs":         p.get("runs", []),
                "run_count":    p.get("run_count", 1),
                "has_bold":     p.get("has_bold", False),
                "partial_bold": p.get("partial_bold", False),
                "has_numpr":    p.get("has_numpr", False),
            }
            for p in content_paras
        ]

        if section_name in sections:
            # Merge into existing entry
            sections[section_name]["raw_text"]      += "\n" + new_text
            sections[section_name]["para_indices"]  += new_indices
            sections[section_name]["lines_format"]  += new_formats
        else:
            sections[section_name] = {
                "raw_text":     new_text,
                "para_indices": new_indices,
                "lines_format": new_formats,
                "heading_index": start_idx,
            }

    # ── Capture specialty tagline in the skills section ──────────────────────
    # Many Canadian resume templates have a "specialty tagline" right after
    # the name/contact block: "Operations & Workforce Coordination · Process
    # Improvement · Administrative Operations". This line is NOT a section
    # heading so _match_section ignores it, and the first-3-skip heuristic
    # below buries it. We detect it by its "A · B · C" or "A | B | C" pattern
    # and inject it into the skills section so the LLM can replace it.
    if section_starts:
        first_section_idx = section_starts[0][1]
        for p in paragraphs:
            if p["index"] >= first_section_idx:
                break
            text = p["text"]
            # Specialty tagline heuristic: contains " · " or " | " separator,
            # all segments are short (< 35 chars), no @ or http (contact lines)
            if ("·" in text or (" | " in text and "@" not in text)) and "http" not in text.lower():
                segments = [s.strip() for s in text.replace("·", "|").split("|")]
                if len(segments) >= 2 and all(len(s) < 45 for s in segments if s):
                    fmt = {
                        "text":         text,
                        "runs":         p.get("runs", []),
                        "run_count":    p.get("run_count", 1),
                        "has_bold":     p.get("has_bold", False),
                        "partial_bold": p.get("partial_bold", False),
                        "has_numpr":    p.get("has_numpr", False),
                    }
                    if "skills" in sections:
                        # Prepend to existing skills section
                        sections["skills"]["raw_text"]     = text + "\n" + sections["skills"]["raw_text"]
                        sections["skills"]["para_indices"] = [p["index"]] + sections["skills"]["para_indices"]
                        sections["skills"]["lines_format"] = [fmt] + sections["skills"]["lines_format"]
                    else:
                        sections["skills"] = {
                            "raw_text":     text,
                            "para_indices": [p["index"]],
                            "lines_format": [fmt],
                            "heading_index": p["index"],
                        }
                    break   # only capture the first matching tagline

    # ── Fallback: capture pre-section text as "summary" ──────────────────────
    # Many Canadian resumes have an unlabelled profile blurb at the top
    # (after name/contact) before the first explicit section heading.
    # Also handles the case where the heading was detected but content is empty.
    summary_empty = "summary" not in sections or not sections["summary"]["raw_text"].strip()
    if summary_empty and section_starts:
        first_section_idx = section_starts[0][1]
        pre_paras = [
            p for p in paragraphs
            if p["index"] < first_section_idx and p["text"] and not _looks_like_heading(p["text"])
        ]
        # Skip the first 3 lines (typically: name, phone/email, location)
        pre_content = pre_paras[3:]
        if pre_content:
            sections["summary"] = {
                "raw_text":     "\n".join(p["text"] for p in pre_content),
                "para_indices": [p["index"] for p in pre_content],
                "lines_format": [
                    {
                        "text":         p["text"],
                        "runs":         p.get("runs", []),
                        "run_count":    p.get("run_count", 1),
                        "has_bold":     p.get("has_bold", False),
                        "partial_bold": p.get("partial_bold", False),
                        "has_numpr":    p.get("has_numpr", False),
                    }
                    for p in pre_content
                ],
                "heading_index": -1,
            }

    return {
        "full_text":      full_text,
        "candidate_name": candidate_name,
        "sections":       sections,
    }


def _match_section(text: str) -> str | None:
    """
    Return the canonical section name if `text` looks like a section heading,
    or None if it doesn't.

    Two safeguards against false positives:
    1. Length cap — real headings are short (≤ 60 chars).
       A paragraph like "Experienced IT systems and technology professional..."
       is content, not a heading, even though it contains the word "experience".
    2. Word boundaries (\b) in the patterns so "experience" doesn't match
       the adjective "Experienced", "skills" doesn't match "skillset", etc.
    """
    stripped = text.strip()
    # Content paragraphs are long; section headings are short
    if len(stripped) > 60:
        return None
    clean = stripped.lower()
    for name, pattern in SECTION_PATTERNS.items():
        if re.search(pattern, clean):
            return name
    return None


def _detect_name(paragraphs: list[dict]) -> str:
    for p in paragraphs[:5]:
        text = p["text"].strip()
        if not text:
            continue
        # Skip lines that look like contact info or headings
        if "@" in text or "http" in text or text.isupper():
            continue
        words = text.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w):
            return text
    return ""

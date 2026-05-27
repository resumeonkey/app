"""
Reconstruct the adapted .docx by modifying ONLY the changed sections
in the original document, preserving all formatting, styles, and layout.

Strategy:
1. Copy the original docx to a new file.
2. Resolve and sort all blocks by their first para_index DESCENDING (last → first).
   This is critical: processing end-to-beginning means any insertion or deletion
   in section N only shifts paragraphs AFTER N, which we have already processed.
   Forward processing causes a cascade where summary's extra line shifts skills into
   a heading slot, which shifts experience into a heading slot, etc.
3. For each section, find its paragraphs by index and update them in-place.
4. When adapted text has more/fewer lines than original, expand/contract carefully.
5. Never touch paragraphs outside changed sections.
"""
import os
import re
import shutil
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from copy import deepcopy
from typing import Any

# Pattern: "| May 2020 – Apr 2021" — marks a job-title line.
# IMPORTANT: must require a MONTH NAME so that certification lines like
# "CertiProf | 2025" or "Udemy | 2024" are NOT mistakenly treated as job titles.
_JOB_TITLE_RE = re.compile(
    r'\|\s*(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|'
    r'jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)',
    re.IGNORECASE,
)

# Sections whose bullet items are intentionally bold (certifications list in
# Canadian resumes).  Bold is always preserved for these, regardless of numPr.
_BOLD_PRESERVE_SECTIONS = frozenset({"certifications"})


def build_adapted_docx(
    master_path: str,
    master_sections: dict[str, Any],
    blocks_changed: list[dict],
    output_path: str,
) -> str:
    """
    Creates output_path as an adapted copy of master_path.
    Returns the output_path.
    """
    # Work on a copy — never mutate the master
    shutil.copy2(master_path, output_path)
    doc = Document(output_path)

    # ── Resolve section info for all blocks ────────────────────────────────────
    resolved: list[tuple[str, str, list[int]]] = []   # (section_name, text, indices)
    for block in blocks_changed:
        section_name = block["section"]
        adapted_text = block["adapted"]

        # Look up by exact key first; fall back to case-insensitive partial match.
        section_info = master_sections.get(section_name)
        if section_info is None:
            for key, info in master_sections.items():
                if section_name.lower() in key.lower() or key.lower() in section_name.lower():
                    section_info = info
                    break

        if not section_info:
            continue

        para_indices = section_info.get("para_indices", [])
        if not para_indices:
            continue

        resolved.append((section_name, adapted_text, para_indices))

    # ── CRITICAL: process from LAST section to FIRST ───────────────────────────
    # Any insertion/deletion in section X only shifts paragraphs at indices > X.
    # Processing in reverse order means those shifts only affect sections we have
    # already processed — so their para_indices stay correct.
    resolved.sort(key=lambda x: x[2][0], reverse=True)

    for section_name, adapted_text, para_indices in resolved:
        _replace_section_content(doc, para_indices, adapted_text, section_name=section_name)

    doc.save(output_path)
    return output_path


def _replace_section_content(
    doc: Document,
    para_indices: list[int],
    new_text: str,
    section_name: str = "",
):
    """
    Replace the text in the paragraphs at para_indices with new_text lines.
    Preserves the formatting (bold, italic, font) of the first run of each paragraph.
    Uses doc.paragraphs (body-level only) — must match resume_parser.parse_docx ordering.
    """
    new_lines = [line for line in new_text.split("\n") if line.strip()]
    paras     = doc.paragraphs

    # Validate indices
    valid_indices = [i for i in para_indices if i < len(paras)]
    if not valid_indices:
        return

    # Safety cap: LLM sometimes expands content far beyond the original.
    # Allow at most 50% more lines than the original slot count to prevent
    # the adapted section from overrunning into adjacent sections.
    max_lines = len(valid_indices) + max(5, len(valid_indices) // 2)
    if len(new_lines) > max_lines:
        new_lines = new_lines[:max_lines]

    paragraphs_to_remove: list = []

    for slot_num, para_idx in enumerate(valid_indices):
        para = paras[para_idx]
        if slot_num < len(new_lines):
            _set_paragraph_text(para, new_lines[slot_num], section_name=section_name)
        else:
            # More original slots than new lines — remove excess paragraphs
            # so they don't leave blank lines in the document.
            paragraphs_to_remove.append(para)

    for para in paragraphs_to_remove:
        _remove_paragraph(para)

    # If more new lines than original paragraphs, insert after the last valid index
    if len(new_lines) > len(valid_indices):
        last_para = paras[valid_indices[-1]]
        extra_lines = new_lines[len(valid_indices):]
        for line in reversed(extra_lines):
            new_para = _insert_paragraph_after(last_para, line)
            _copy_style(last_para, new_para)


def _remove_paragraph(para) -> None:
    """Physically remove a paragraph from the document XML."""
    p = para._p
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _set_paragraph_text(para, text: str, section_name: str = ""):
    """
    Clear all runs in paragraph and set new text, preserving first-run formatting.

    Job-title detection:
      Lines matching _JOB_TITLE_RE (contains | + month name) are forced to bold
      headings with numPr removed, regardless of the original paragraph style.

    Bold logic for non-title lines:
      • Non-bullet paragraphs (numPr=False): bold is always preserved.
        These are heading/title slots — their bold is intentional.
      • Bullet paragraphs in certifications: bold is preserved
        (certifications are intentionally bold bullets in Canadian resumes).
      • Bullet paragraphs in all other sections: bold is stripped as a safety
        net against overflow content inheriting bold from a wrong slot.
        Since experience bullets are not bold in the master, stripping is
        usually a no-op; it only matters in rare overflow edge cases.
    """
    from docx.oxml import OxmlElement

    # Strip leading bullet chars that some LLMs prepend
    clean_text = text.lstrip("•-– ").strip() if text.startswith(("•", "-", "–")) else text
    if not clean_text:
        for run in para.runs:
            run._r.getparent().remove(run._r)
        for r in para._p.findall(qn("w:r")):
            para._p.remove(r)
        return

    is_job_title = bool(_JOB_TITLE_RE.search(clean_text))

    if is_job_title:
        # ── Job-title line: force bold heading, remove bullet properties ──────
        for run in para.runs:
            run._r.getparent().remove(run._r)
        for r in para._p.findall(qn("w:r")):
            para._p.remove(r)

        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None:
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                pPr.remove(numPr)

        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rPr.append(b)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = clean_text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        para._p.append(r)

    else:
        # ── Normal line: conditional bold preservation ────────────────────────
        first_run_rpr = None
        if para.runs:
            rpr = para.runs[0]._r.find(qn("w:rPr"))
            if rpr is not None:
                first_run_rpr = deepcopy(rpr)

                # Determine whether to strip bold:
                # • Non-bullet slots (numPr=False): preserve bold always —
                #   these are title/heading slots, bold is intentional.
                # • Certifications bullets: preserve bold (intentional).
                # • All other bullet slots: strip bold as safety net.
                orig_has_numpr = (
                    para._p.find(qn("w:pPr")) is not None
                    and para._p.find(qn("w:pPr")).find(qn("w:numPr")) is not None
                )
                should_strip_bold = (
                    orig_has_numpr
                    and section_name not in _BOLD_PRESERVE_SECTIONS
                )
                if should_strip_bold:
                    for bold_tag in ("w:b", "w:bCs"):
                        el = first_run_rpr.find(qn(bold_tag))
                        if el is not None:
                            first_run_rpr.remove(el)

        for run in para.runs:
            run._r.getparent().remove(run._r)
        for r in para._p.findall(qn("w:r")):
            para._p.remove(r)

        r = OxmlElement("w:r")
        if first_run_rpr is not None:
            r.append(deepcopy(first_run_rpr))
        t = OxmlElement("w:t")
        t.text = clean_text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        para._p.append(r)


def _insert_paragraph_after(ref_para, text: str):
    """Insert a new paragraph after ref_para with the same style."""
    from docx.oxml import OxmlElement
    new_p = OxmlElement("w:p")

    ref_ppr = ref_para._p.find(qn("w:pPr"))
    if ref_ppr is not None:
        new_p.append(deepcopy(ref_ppr))

    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    new_p.append(r)

    ref_para._p.addnext(new_p)

    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, ref_para._p.getparent())


def _copy_style(source_para, target_para):
    """Copy paragraph style name from source to target."""
    try:
        if source_para.style:
            target_para.style = source_para.style
    except Exception:
        pass

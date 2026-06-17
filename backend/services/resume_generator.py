"""
Generate a clean DOCX resume from structured data (NOT by editing an uploaded
file). Because we control the whole document structure, the output formatting is
100% predictable — this eliminates the parsing/in-place-edit bugs (broken tables,
lost bullet colours, leaked debug lines, mangled sections).

Public API:
    generate_resume_docx(data: dict, out_path: str, template: str = "classic") -> str
"""
from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Templates (colour + font tokens) ─────────────────────────────────────────
_TEMPLATES: dict[str, dict] = {
    "classic": {  # the blue look of the current master
        "name_color":    "1F4E79",
        "title_color":   "2E75B6",
        "heading_color": "1F4E79",
        "accent_color":  "1F4E79",
        "text_color":    "1A1A1A",
        "muted_color":   "444444",
        "font":          "Calibri",
        "name_size":     22,
        "title_size":    13,
    },
    "iris": {  # SoyManada "Iris & Ivory"
        "name_color":    "3D1A78",
        "title_color":   "7B4DC8",
        "heading_color": "5B2D9E",
        "accent_color":  "7B4DC8",
        "text_color":    "120826",
        "muted_color":   "5A4877",
        "font":          "Calibri",
        "name_size":     22,
        "title_size":    13,
    },
}


def _rgb(hexstr: str) -> RGBColor:
    return RGBColor.from_string(hexstr)


def generate_resume_docx(data: dict, out_path: str, template: str = "classic") -> str:
    t = _TEMPLATES.get(template, _TEMPLATES["classic"])
    doc = Document()

    # Tighter margins for a one/two-page resume
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.5)
        s.left_margin = s.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = t["font"]
    normal.font.size = Pt(10)
    normal.font.color.rgb = _rgb(t["text_color"])

    def _spacing(p, before=0, after=2, line=1.0):
        pf = p.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = line

    def _run(p, text, *, bold=False, size=10, color=None):
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = _rgb(color)
        return r

    def _heading(text):
        p = doc.add_paragraph()
        _spacing(p, before=8, after=3)
        _run(p, text.upper(), bold=True, size=11, color=t["heading_color"])
        # bottom border under the heading
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), t["heading_color"])
        pbdr.append(bottom)
        pPr.append(pbdr)
        return p

    # ── Header: name / title / tagline / contact ─────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _spacing(p, after=0)
    _run(p, data.get("name", ""), bold=True, size=t["name_size"], color=t["name_color"])

    if data.get("title"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _spacing(p, after=1)
        _run(p, data["title"], bold=True, size=t["title_size"], color=t["title_color"])

    if data.get("tagline"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _spacing(p, after=1)
        _run(p, "  ·  ".join(data["tagline"]), size=9, color=t["muted_color"])

    c = data.get("contact", {})
    contact_bits = [c.get("location"), c.get("phone"), c.get("email"), c.get("linkedin"), c.get("website")]
    contact_bits = [b for b in contact_bits if b]
    if contact_bits:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _spacing(p, after=4)
        _run(p, "  •  ".join(contact_bits), size=9, color=t["muted_color"])

    # ── Professional summary ─────────────────────────────────────────────────
    if data.get("summary"):
        _heading("Professional Summary")
        for para in data["summary"]:
            p = doc.add_paragraph(); _spacing(p, after=3, line=1.05)
            _run(p, para, size=10, color=t["muted_color"])

    # ── Core competencies (3-column table) ───────────────────────────────────
    comps = data.get("competencies") or {}
    if comps:
        _heading("Core Competencies")
        cols = list(comps.keys())
        table = doc.add_table(rows=1, cols=len(cols))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for ci, cat in enumerate(cols):
            cell = table.rows[0].cells[ci]
            cell.paragraphs[0].text = ""
            hp = cell.paragraphs[0]; _spacing(hp, after=2)
            _run(hp, cat, bold=True, size=10, color=t["accent_color"])
            for skill in comps[cat]:
                bp = cell.add_paragraph(); _spacing(bp, after=0, line=1.0)
                _run(bp, "•  ", bold=True, color=t["accent_color"])
                _run(bp, skill, size=9.5, color=t["text_color"])

    # ── Professional experience ──────────────────────────────────────────────
    if data.get("experience"):
        _heading("Professional Experience")
        for job in data["experience"]:
            p = doc.add_paragraph(); _spacing(p, before=4, after=0)
            _run(p, job.get("title", ""), bold=True, size=10.5, color=t["text_color"])
            loc = " – ".join(x for x in [job.get("company"), job.get("location")] if x)
            if loc:
                _run(p, "  –  " + loc, size=10, color=t["muted_color"])
            if job.get("dates"):
                _run(p, "\t" + job["dates"], size=9.5, color=t["muted_color"])
            for b in job.get("bullets", []):
                bp = doc.add_paragraph(); _spacing(bp, after=1, line=1.05)
                bp.paragraph_format.left_indent = Inches(0.2)
                _run(bp, "•  ", bold=True, color=t["accent_color"])
                _run(bp, b, size=10, color=t["text_color"])

    # ── Education ────────────────────────────────────────────────────────────
    if data.get("education"):
        _heading("Education")
        for ed in data["education"]:
            p = doc.add_paragraph(); _spacing(p, after=1)
            _run(p, ed.get("degree", ""), bold=True, size=10, color=t["text_color"])
            if ed.get("institution"):
                _run(p, "  –  " + ed["institution"], size=9.5, color=t["muted_color"])
            if ed.get("dates"):
                _run(p, "\t" + ed["dates"], size=9.5, color=t["muted_color"])

    # ── Certifications ───────────────────────────────────────────────────────
    if data.get("certifications"):
        _heading("Certifications")
        for ce in data["certifications"]:
            p = doc.add_paragraph(); _spacing(p, after=1)
            _run(p, ce.get("name", ""), bold=True, size=10, color=t["text_color"])
            tail = " – ".join(x for x in [ce.get("issuer"), ce.get("year")] if x)
            if tail:
                _run(p, "  –  " + tail, size=9.5, color=t["muted_color"])

    doc.save(out_path)
    return out_path
